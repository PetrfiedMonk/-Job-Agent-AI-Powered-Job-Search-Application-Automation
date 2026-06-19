"""
Resume Builder
Converts a TailoredResume into a polished DOCX file.
Uses a clean, ATS-friendly template optimized for interview conversion.
"""
import os
from pathlib import Path
from datetime import datetime
from typing import Optional

from job_agent.models import TailoredResume, UserProfile


def build_resume_docx(
    resume: TailoredResume,
    output_dir: str,
    template: str = "modern"
) -> str:
    """
    Generate a DOCX resume from a TailoredResume object.
    Returns the path to the generated file.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise ImportError(
            "python-docx is required for resume generation.\n"
            "Run: pip install python-docx"
        )

    os.makedirs(output_dir, exist_ok=True)

    profile = resume.profile
    job = resume.job

    # Clean filename: "Justin_Carano_ProductManager_Acme.docx"
    safe_title = "".join(c for c in job.title if c.isalnum() or c in " _-")[:30].strip().replace(" ", "_")
    safe_company = "".join(c for c in job.company if c.isalnum() or c in " _-")[:20].strip().replace(" ", "_")
    name_part = profile.name.replace(" ", "_") if profile.name else "Resume"
    filename = f"{name_part}_{safe_title}_{safe_company}.docx"
    filepath = Path(output_dir) / filename

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    # ── Default paragraph style ───────────────────────────────────────────────
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)

    # ══════════ HEADER ════════════════════════════════════════════════════════
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(profile.name or "Your Name")
    name_run.bold = True
    name_run.font.size = Pt(22)
    name_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # Contact line
    contact_parts = [x for x in [
        profile.phone, profile.email, profile.location,
        profile.linkedin_url, profile.website
    ] if x]
    contact_para = doc.add_paragraph(" | ".join(contact_parts))
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.runs[0].font.size = Pt(9.5)
    contact_para.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    _add_divider(doc)

    # ══════════ SUMMARY ═══════════════════════════════════════════════════════
    _section_heading(doc, "PROFESSIONAL SUMMARY")
    summary_para = doc.add_paragraph(resume.tailored_summary or profile.summary)
    summary_para.paragraph_format.space_after = Pt(6)

    # ══════════ EXPERIENCE ════════════════════════════════════════════════════
    _section_heading(doc, "EXPERIENCE")
    for exp in resume.tailored_experience:
        # Title + dates on same line
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        title_run = p.add_run(exp.title)
        title_run.bold = True
        title_run.font.size = Pt(11)

        dates = f"  {exp.start_date or ''} – {exp.end_date or 'Present'}"
        date_run = p.add_run(dates)
        date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        date_run.font.size = Pt(10)

        # Company
        company_para = doc.add_paragraph(exp.company)
        company_para.paragraph_format.space_after = Pt(2)
        company_para.runs[0].italic = True
        company_para.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x88)

        # Description (if any)
        if exp.description:
            desc_para = doc.add_paragraph(exp.description)
            desc_para.paragraph_format.space_after = Pt(2)

        # Achievement bullets
        for achievement in exp.achievements:
            bullet = doc.add_paragraph(style="List Bullet")
            bullet.paragraph_format.left_indent = Inches(0.2)
            bullet.paragraph_format.space_after = Pt(1)
            bullet.add_run(achievement)

    # ══════════ SKILLS ════════════════════════════════════════════════════════
    _section_heading(doc, "SKILLS")
    skills_text = " • ".join(resume.highlighted_skills or profile.skills[:20])
    skills_para = doc.add_paragraph(skills_text)
    skills_para.paragraph_format.space_after = Pt(4)

    # ══════════ CERTIFICATIONS ════════════════════════════════════════════════
    if profile.certifications:
        _section_heading(doc, "CERTIFICATIONS")
        cert_para = doc.add_paragraph(" • ".join(profile.certifications))
        cert_para.paragraph_format.space_after = Pt(4)

    # ══════════ PROJECTS ══════════════════════════════════════════════════════
    if profile.projects:
        _section_heading(doc, "PROJECTS")
        for proj in profile.projects[:3]:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            proj_run = p.add_run(proj.get("name", ""))
            proj_run.bold = True
            p.add_run(f"  {proj.get('description', '')}")
            if proj.get("impact"):
                p.add_run(f" | {proj['impact']}")

    # ══════════ EDUCATION ═════════════════════════════════════════════════════
    _section_heading(doc, "EDUCATION")
    for edu in profile.education:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        deg_run = p.add_run(f"{edu.degree}")
        deg_run.bold = True
        p.add_run(f"  {edu.school}")
        if edu.year:
            p.add_run(f"  |  {edu.year}")

    doc.save(str(filepath))
    resume.docx_path = str(filepath)
    print(f"[builder] Resume saved: {filepath.name}")
    return str(filepath)


def _section_heading(doc, text: str):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    # Bottom border
    _add_para_bottom_border(p)


def _add_divider(doc):
    from docx.shared import Pt
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    _add_para_bottom_border(p)


def _add_para_bottom_border(para):
    """Add a bottom border line to a paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1A1A2E")
    pBdr.append(bottom)
    pPr.append(pBdr)
