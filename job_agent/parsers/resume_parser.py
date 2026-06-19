"""
Resume Parser
Extracts structured data from PDF or DOCX resume files.
"""
import re
from pathlib import Path
from typing import Dict, List, Optional


def parse_resume(path: str) -> Dict:
    """
    Parse a resume file (PDF or DOCX) and return raw text + structured sections.

    Returns:
        {
            "raw_text": str,
            "sections": {section_name: content},
            "contact": {...},
        }
    """
    filepath = Path(path)
    if not filepath.exists():
        raise FileNotFoundError(f"Resume not found: {path}")

    ext = filepath.suffix.lower()

    if ext == ".pdf":
        raw_text = _parse_pdf(filepath)
    elif ext in (".docx", ".doc"):
        raw_text = _parse_docx(filepath)
    elif ext in (".txt", ".md"):
        raw_text = filepath.read_text(encoding="utf-8", errors="replace")
    else:
        raise ValueError(f"Unsupported resume format: {ext}. Use PDF, DOCX, or TXT.")

    sections = _extract_sections(raw_text)
    contact = _extract_contact(raw_text)

    return {
        "raw_text": raw_text,
        "sections": sections,
        "contact": contact,
        "file_path": str(filepath),
        "file_type": ext,
    }


def _parse_pdf(filepath: Path) -> str:
    """Extract text from PDF using pdfminer or pypdf2."""
    try:
        import pdfminer.high_level
        return pdfminer.high_level.extract_text(str(filepath))
    except ImportError:
        pass

    try:
        import pypdf
        reader = pypdf.PdfReader(str(filepath))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)
    except ImportError:
        pass

    raise ImportError(
        "PDF parsing requires pdfminer.six or pypdf.\n"
        "Run: pip install pdfminer.six"
    )


def _parse_docx(filepath: Path) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(str(filepath))
        paragraphs = [p.text for p in doc.paragraphs]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                paragraphs.extend([cell.text for cell in row.cells])
        return "\n".join(paragraphs)
    except ImportError:
        raise ImportError(
            "DOCX parsing requires python-docx.\n"
            "Run: pip install python-docx"
        )


# Section headers commonly found in resumes
SECTION_PATTERNS = {
    "summary": r"(summary|objective|profile|about me|professional summary)",
    "experience": r"(experience|work history|employment|work experience|professional experience)",
    "education": r"(education|academic|degree|university|college|school)",
    "skills": r"(skills|technical skills|core competencies|technologies|expertise)",
    "certifications": r"(certif|license|credential|accreditation)",
    "projects": r"(project|portfolio|work samples|open source)",
    "achievements": r"(achievement|award|honor|recognition|accomplishment)",
    "volunteer": r"(volunteer|community|non-?profit)",
}


def _extract_sections(text: str) -> Dict[str, str]:
    """Split resume text into named sections."""
    lines = text.split("\n")
    sections = {}
    current_section = "header"
    current_lines = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            current_lines.append("")
            continue

        # Check if this line is a section header
        matched_section = None
        for section_name, pattern in SECTION_PATTERNS.items():
            if re.match(rf"^{pattern}[\s:]*$", stripped, re.IGNORECASE):
                matched_section = section_name
                break

        if matched_section:
            # Save previous section
            sections[current_section] = "\n".join(current_lines).strip()
            current_section = matched_section
            current_lines = []
        else:
            current_lines.append(line)

    # Save last section
    sections[current_section] = "\n".join(current_lines).strip()
    return {k: v for k, v in sections.items() if v}


def _extract_contact(text: str) -> Dict[str, str]:
    """Extract contact info from resume text."""
    contact = {}

    # Email
    email_match = re.search(r'[\w.+-]+@[\w-]+\.[a-zA-Z]{2,}', text)
    if email_match:
        contact["email"] = email_match.group()

    # Phone
    phone_match = re.search(
        r'(?:\+?1[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}', text
    )
    if phone_match:
        contact["phone"] = phone_match.group().strip()

    # LinkedIn
    linkedin_match = re.search(r'linkedin\.com/in/[\w-]+', text, re.IGNORECASE)
    if linkedin_match:
        contact["linkedin"] = "https://" + linkedin_match.group()

    # GitHub
    github_match = re.search(r'github\.com/[\w-]+', text, re.IGNORECASE)
    if github_match:
        contact["github"] = "https://" + github_match.group()

    # Website
    website_match = re.search(
        r'https?://(?!linkedin|github)[\w.-]+\.[a-zA-Z]{2,}(?:/[\w.-]*)?',
        text, re.IGNORECASE
    )
    if website_match:
        contact["website"] = website_match.group()

    return contact
